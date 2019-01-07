from PyQt5 import QtGui
from PyQt5.QtGui import QIcon
from PyQt5.QtWidgets import QAction, QTextEdit, QColorDialog, QFontDialog, QApplication, QMainWindow, QPushButton, \
    QMessageBox, QMenu, QAction, QMenuBar, QStatusBar
from PyQt5.QtPrintSupport import QPrintDialog, QPrinter, QPrintPreviewDialog
import subprocess
# بداية برنامج الارسال
from openpyxl.styles import Font
from openpyxl.styles.colors import Color
import docx, openpyxl, os, sys, os.path, string
from docx.shared import RGBColor
import comtypes.client, time
import httplib2
import oauth2client
from oauth2client import client, tools
import base64
from openpyxl.styles import Alignment
import mimetypes
from email import encoders
from email.message import Message
from email.mime.audio import MIMEAudio
from email.mime.base import MIMEBase
from email.mime.image import MIMEImage
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from apiclient import errors, discovery
from docx.shared import Inches

SCOPES = 'https://www.googleapis.com/auth/gmail.send'
CLIENT_SECRET_FILE = 'client_secret_social.json'
APPLICATION_NAME = 'Gmail API Python Send Email'


def get_credentials():
    home_dir = os.path.expanduser('~')  # >> C:\Users\Me
    credential_dir = os.path.join(home_dir, '.credentials')  # >>C:\Users\Me\.credentials   (it's a folder)
    if not os.path.exists(credential_dir):
        os.makedirs(credential_dir)
    credential_path = os.path.join(credential_dir, 'cred send mail.json')
    store = oauth2client.file.Storage(credential_path)
    credentials = store.get()

    if not credentials or credentials.invalid:
        # CLIENT_SECRET_FILE = 'client_secret_social.json'
        # APPLICATION_NAME = 'Gmail API Python Send Email'
        # SCOPES = 'https://www.googleapis.com/auth/gmail.send'
        flow = client.flow_from_clientsecrets(CLIENT_SECRET_FILE, SCOPES)
        flow.user_agent = APPLICATION_NAME
        credentials = tools.run_flow(flow, store)

    return credentials


def create_message_and_send(sender, to, subject, message_text_plain, message_text_html, attached_file):
    credentials = get_credentials()

    # Create an httplib2.Http object to handle our HTTP requests, and authorize it using credentials.authorize()
    http = httplib2.Http()

    # http is the authorized httplib2.Http()
    http = credentials.authorize(http)  # or: http = credentials.authorize(httplib2.Http())

    service = discovery.build('gmail', 'v1', http=http)

    ## without attachment
    # message_without_attachment = create_message_without_attachment(sender, to, subject, message_text_html, message_text_plain)
    # send_Message_without_attachement(service, "me", message_without_attachment, message_text_plain)

    ## with attachment
    message_with_attachment = create_Message_with_attachment(sender, to, subject, message_text_plain, message_text_html,
                                                             attached_file)
    send_Message_with_attachement(service, "me", message_with_attachment)


def create_message_without_attachment(sender, to, subject, message_text_html, message_text_plain):
    # Create message container
    message = MIMEMultipart('alternative')  # needed for both plain & HTML (the MIME type is multipart/alternative)
    message['Subject'] = subject
    message['From'] = sender
    message['To'] = to

    # Create the body of the message (a plain-text and an HTML version)
    message.attach(MIMEText(message_text_plain, 'plain'))
    message.attach(MIMEText(message_text_html, 'html'))

    raw_message_no_attachment = base64.urlsafe_b64encode(message.as_bytes())
    raw_message_no_attachment = raw_message_no_attachment.decode()
    body = {'raw': raw_message_no_attachment}
    return body


def create_Message_with_attachment(sender, to, subject, message_text_plain, message_text_html, attached_file):
    message = MIMEMultipart()
    message['to'] = to
    message['from'] = sender
    message['subject'] = subject

    message.attach(MIMEText(message_text_html, 'html'))
    message.attach(MIMEText(message_text_plain, 'plain'))

    my_mimetype, encoding = mimetypes.guess_type(attached_file)

    if my_mimetype is None or encoding is not None:
        my_mimetype = 'application/octet-stream'

    main_type, sub_type = my_mimetype.split('/', 1)
    if main_type == 'text':
        print("text")
        temp = open(attached_file, 'r')
        attachement = MIMEText(temp.read(), _subtype=sub_type)
        temp.close()

    elif main_type == 'image':
        print("image")
        temp = open(attached_file, 'rb')
        attachement = MIMEImage(temp.read(), _subtype=sub_type)
        temp.close()

    elif main_type == 'audio':
        print("audio")
        temp = open(attached_file, 'rb')
        attachement = MIMEAudio(temp.read(), _subtype=sub_type)
        temp.close()

    elif main_type == 'application' and sub_type == 'pdf':
        temp = open(attached_file, 'rb')
        attachement = MIMEApplication(temp.read(), _subtype=sub_type)
        temp.close()

    else:
        attachement = MIMEBase(main_type, sub_type)
        temp = open(attached_file, 'rb')
        attachement.set_payload(temp.read())
        temp.close()

    encoders.encode_base64(attachement)  # https://docs.python.org/3/library/email-examples.html
    filename = os.path.basename(attached_file)
    attachement.add_header('Content-Disposition', 'attachment', filename=filename)  # name preview in email
    message.attach(attachement)
    message_as_bytes = message.as_bytes()  # the message should converted from string to bytes.
    message_as_base64 = base64.urlsafe_b64encode(message_as_bytes)  # encode in base64 (printable letters coding)
    raw = message_as_base64.decode()  # need to JSON serializable (no idea what does it means)
    return {'raw': raw}


def send_Message_without_attachement(service, user_id, body, message_text_plain):
    try:
        message_sent = (service.users().messages().send(userId=user_id, body=body).execute())
        message_id = message_sent['id']
        print(f'Message sent (without attachment) \n\n Message Id: {message_id}\n\n Message:\n\n {message_text_plain}')
    except errors.HttpError as error:
        print(f'An error occurred: {error}')


def send_Message_with_attachement(service, user_id, message_with_attachment):
    try:
        message_sent = (service.users().messages().send(userId=user_id, body=message_with_attachment).execute())
        message_id = message_sent['id']
    except errors.HttpError as error:
        print(f'An error occurred: {error}')


# end imported to send email


#  نهاية متطلبات برنامج الارسال  الارسال

def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)


def totalPdf(encrypt_pdf_path, basePath, fileName):
    os.chdir(encrypt_pdf_path)
    cmd = 'pdftk *.pdf cat output Total_iftar_{}.pdf'.format(fileName)
    os.system(cmd)
    os.chdir(basePath)


def sendEmail_spec():
    window.textEdit.clear()
    print('sending email spec. Letter ')
    basePath = os.getcwd()
    Excl_file = "excel_Data/iftar_spec.xlsx"  # ملف اكسل يحتوي على معلومات كل الجهات المراد ارسالها
    encrypt_pdf_path = basePath + "\\iftar_spec"
    wb = openpyxl.load_workbook(Excl_file)
    sheet = wb['Data']
    NumRecord = sheet.max_row  # تحديد عدد الاسطر في الجدول
    print('Number of Records : ', NumRecord)
    r = 3
    while (r < NumRecord):
        # centerName = str(sheet.cell(row=r, column=3).value)
        assosiation = sheet.cell(row=r, column=5)
        receiverName = sheet.cell(row=r, column=9)
        Email = sheet.cell(row=r, column=7)
        outNum = sheet.cell(row=r, column=15)
        flag = sheet.cell(row=r, column=17).value
        status = sheet.cell(row=r, column=18)

        # to = 'r_maher@wamy.org'
        # Email_To_Send =[Email]
        # to = 'amedshel@wamy.org'
        to = Email.value

        encrypted_file = os.path.join(encrypt_pdf_path, r"{}.pdf".format(outNum.value))
        print(encrypted_file)
        # جزء البرنامج الذي يقوم بارسال بريد الكتروني لكل شركة
        # fo Email send information
        sender = "social@wamy.org"
        subject = 'تعميد الجهات المشرفة علي إفطار الصائم 1439هـ‎'
        message_text_plain = ""
        attached_file = encrypted_file
        message_text_html = """<blockquote>
                <p><b><div align="center" ><h2>   الاخوة الأفاضل  : {}</h2><br>
                 آمل ان يصلكم خطابنا هذا وأنتم في صحة طيبة وعلى خير ما يحب ربنا ويرضى   <br>  
                   مرفق لكم تعميدكم بعدد وجبات  إفطار الصائم لعام 1440 هـ   <br>

                   نأمل التواصل مع ( <font size="4" color="blue">{}</font>) لتنفيذ البرنامج  <br><br>

                 <b><h2>الندوة العالمية للشباب الاسلامي<br>
                   إدارة الشؤون الإجتماعية<br> 
                          </b></h2>
                </div></blockquote>""".format(assosiation.value, receiverName.value)
        if flag == 1:
            print('flag = 1')
            if to != None:
                print("now send Email to : {}".format(to))
                create_message_and_send(sender, to, subject, message_text_plain, message_text_html, attached_file)
                status.alignment = Alignment(horizontal='center', vertical='center')
                status.font = Font(color="FF0000")
                status.value = " تم الارسال بنجاح"
                window.done(to, assosiation.value, outNum.value)
                time.sleep(0.2)
            else:
                print("لا يوجد بريد لارساله ")
                status.value = "لا يوجد عنوان بريد "
                time.sleep(0.2)
        else:
            print('Flag = 0')
        r += 1
    wb.save(Excl_file)
    print('done......')


def sendLetter():
    window.textEdit.clear()
    print('sending email spec. Letter ')
    basePath = os.getcwd()
    Excl_file = "excel_Data/iftar.xlsx"  # ملف اكسل يحتوي على معلومات كل الجهات المراد ارسالها
    encrypt_pdf_path = basePath + "\\iftar"
    wb = openpyxl.load_workbook(Excl_file)
    sheet = wb['Data']
    NumRecord = sheet.max_row  # تحديد عدد الاسطر في الجدول
    print('Number of Records : ', NumRecord)
    r = 5
    while (r < (NumRecord + 1)):
        assosiation = sheet.cell(row=r, column=2)
        receiverName = sheet.cell(row=r, column=9)
        Email = sheet.cell(row=r, column=8)
        outNum = sheet.cell(row=r, column=13)
        flag = sheet.cell(row=r, column=15).value
        status = sheet.cell(row=r, column=16)

        # to = 'r_maher@wamy.org'
        # Email_To_Send =[Email]
        # to = 'amedshel@wamy.org'
        to = Email.value
        print('test')

        encrypted_file = os.path.join(encrypt_pdf_path, r"{}.pdf".format(outNum.value))
        print(encrypted_file)
        # جزء البرنامج الذي يقوم بارسال بريد الكتروني لكل شركة
        # fo Email send information
        sender = "social@wamy.org"
        subject = 'تعميد الجهات المشرفة علي إفطار الصائم 1440هـ'
        message_text_plain = ""
        attached_file = encrypted_file
        message_text_html = """<blockquote>
                   <p><b><div align="center" ><h2>   الاخوة الأفاضل  : {}</h2><br>
                    آمل ان يصلكم خطابنا هذا وأنتم في صحة طيبة وعلى خير ما يحب ربنا ويرضى   <br>  
                      مرفق لكم تعميدكم بعدد وجبات  إفطار الصائم لعام 1440 هـ   <br>

                      نأمل التواصل مع ( <font size="4" color="blue">{}</font>) لتنفيذ البرنامج  <br><br>

                    <b><h2>الندوة العالمية للشباب الاسلامي<br>
                      إدارة الشؤون الإجتماعية<br> 
                             </b></h2>
                   </div></blockquote>""".format(assosiation.value, receiverName.value)
        if flag == 1:
            print('flag = 1')
            if to != None:
                print("now send Email to : {}".format(to))
                create_message_and_send(sender, to, subject, message_text_plain, message_text_html, attached_file)
                status.alignment = Alignment(horizontal='center', vertical='center')
                status.font = Font(color="FF0000")
                status.value = " تم الارسال بنجاح"
                window.done(to, assosiation.value, outNum.value)
                time.sleep(0.2)
            else:
                print("لا يوجد بريد لارساله ")
                status.value = "لا يوجد عنوان بريد "
                time.sleep(0.2)
        else:
            print('Flag = 0')
        r += 1
    wb.save(Excl_file)
    print('done......')


def createLetter():
    """*******************************************************************************************
    **************************     انشاء خطابات غير المخصص ******************************************
    **********************************************************************************************
    """
    basePath = os.getcwd()
    print('this is base path : ', os.getcwd())
    owner_password = "Wamy@12379"
    Excl_file = "excel_Data/iftar.xlsx"  # ملف اكسل يحتوي على معلومات كل الجهات المراد ارسالها
    Word_file = 'word_Data/iftar.docx'  # ملف خطاب المراد ارسالة الخطاب نسخة منه
    outLetter_path = basePath + "\\outLetter"
    outpdf_path = basePath + "\\outpdf"
    encrypt_pdf_path = basePath + "\\iftar"
    if not os.path.exists(outLetter_path):
        print('Creating output folder...')
        os.makedirs(outLetter_path)
        print(outLetter_path, 'created.')
    else:
        print(outLetter_path, 'already exists.\n')

    if not os.path.exists(outpdf_path):
        print('Creating output folder...')
        os.makedirs(outpdf_path)
        print(outpdf_path, 'created.')
    else:
        print(outpdf_path, 'already exists.\n')

    if not os.path.exists(encrypt_pdf_path):
        print('Creating output folder...')
        os.makedirs(encrypt_pdf_path)
        print(encrypt_pdf_path, 'created.')
    else:
        print(encrypt_pdf_path, 'already exists.\n')

    wb = openpyxl.load_workbook(Excl_file)
    sheet = wb['Data']

    doc = docx.Document(Word_file)
    # p = doc.add_paragraph('test')
    NumRecord = sheet.max_row  # تحديد عدد الاسطر في الجدول
    print('Number of Records : ', NumRecord)
    for r in range(5, NumRecord):
        assosiation = sheet.cell(row=r, column=2)
        office = sheet.cell(row=r, column=3).value
        country = sheet.cell(row=r, column=4).value
        receiverName = sheet.cell(row=r, column=9)
        numMeal = sheet.cell(row=r, column=12)
        Email = sheet.cell(row=r, column=8)
        outNum = sheet.cell(row=r, column=13)
        outDate = sheet.cell(row=r, column=14)
        salesName = sheet.cell(row=r, column=10)
        salesMobile = sheet.cell(row=r, column=11)
        flag = sheet.cell(row=r, column=15)
        status = sheet.cell(row=r, column=16)
        copy = sheet.cell(row=r, column=17).value
        to = Email.value
        print('تم توليد الخطاب رقم : ', outNum.value)
        doc.paragraphs[0].runs[3].text = str(outNum.value)
        doc.paragraphs[1].runs[3].text = str(outDate.value)
        doc.paragraphs[6].runs[1].text = assosiation.value
        doc.paragraphs[10].runs[5].text = receiverName.value
        # اضافة صفر واحد على رقم الجوال المقروء من جدول اكسل اذا كانت الصيغة المطلوبة محلية وصفرين اذا كانت الصيغة دولية
        doc.paragraphs[11].runs[3].text = str(salesMobile.value)
        doc.paragraphs[11].runs[1].text = salesName.value
        # doc.paragraphs[6].runs[2].text = ' '*int(55-int(1.10*len(assosiation.value)))
        # print(len(assosiation.value),' : ',int(56-int(1.15*len(assosiation.value))))
        doc.paragraphs[10].runs[7].text = str(numMeal.value)

        if copy == 1:
            # p.runs[0].text = '-صورة مع التحية لمدير مكتب {} '.format(country)
            doc.paragraphs[23].runs[0].text = '-'
            doc.paragraphs[23].runs[1].text = 'صورة مع التحية لمدير مكتب {} '.format(office)
        else:
            # p.runs[0].text = ''
            doc.paragraphs[23].runs[1].text = ''
            doc.paragraphs[23].runs[0].text = ''

        doc_file = os.path.join(outLetter_path, r"{}.docx".format(outNum.value))
        out_file = os.path.join(outpdf_path, r"{}.pdf".format(outNum.value))
        encrypted_file = os.path.join(encrypt_pdf_path, r"{}.pdf".format(outNum.value))

        doc.save(doc_file)
        # p.runs[0].text = ''
        # جزء البرنامج الذي يقوم بتحويل الخطاب من ورد الى بي دي اف
        wdFormatPDF = 17
        in_file = doc_file

        word = comtypes.client.CreateObject('Word.Application')
        docfinal = word.Documents.Open(in_file)
        docfinal.SaveAs(out_file, FileFormat=wdFormatPDF)
        docfinal.Close()
        word.Quit()

        # **************************************
        #    التشفير باستخدام ادوات pdftk
        #    %%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%
        # print(out_file)

        cmd = 'pdftk "{}" input_pw "{}" background background.pdf output "{}" encrypt_128bit allow "printing"'.format(
            out_file, owner_password,
            encrypted_file)

        os.system(cmd)
        # subprocess.call(cmd)
        # %%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%
        # جزء البرنامج الذي يقوم بارسال بريد الكتروني لكل شركة
        # fo Email send information
        sender = "social@wamy.org"
        subject = 'تعميد الجهات المشرفة علي إفطار الصائم 1440هـ'
        message_text_plain = ""
        attached_file = encrypted_file
        message_text_html = """<blockquote>
        <p><b><div align="center" ><h2>   الاخوة الأفاضل  : {}</h2><br>
         آمل ان يصلكم خطابنا هذا وأنتم في صحة طيبة وعلى خير ما يحب ربنا ويرضى   <br>  
           مرفق لكم تعميدكم بعدد وجبات  إفطار الصائم لعام 1440 هـ   <br>

           نأمل التواصل مع ( <font size="4" color="blue">{}</font>) لتنفيذ البرنامج  <br><br>

         <b><h2>الندوة العالمية للشباب الاسلامي<br>
           إدارة الشؤون الإجتماعية<br> 
                  </b></h2>
        </div></blockquote>""".format(assosiation.value, receiverName.value)
        if outNum.value != None:
            # create_message_and_send(sender, to, subject, message_text_plain, message_text_html, attached_file)
            status.alignment = Alignment(horizontal='center', vertical='center')
            status.font = Font(color="FF0000")
            flag.value = 1
            # window.done(to, assosiation.value, outNum.value)
            window.textEdit.append('تم توليد الخطاب رقم :   {} '.format(outNum.value))

        else:
            print("لا يوجد اسم صادر  ")
            flag.value = 0

    wb.save(Excl_file)
    # os.system('taskkill /F /IM WINWORD.EXE')
    totalPdf(encrypt_pdf_path, basePath, sheet.cell(row=5, column=13).value[:4])

    print("done")


def create_Spec():
    window.textEdit.clear()
    basePath = os.getcwd()
    print('this is base path : ', os.getcwd())
    owner_password = "Wamy@12379"
    Excl_file = "excel_Data/iftar_spec.xlsx"  # ملف اكسل يحتوي على معلومات كل الجهات المراد ارسالها
    Word_file = 'word_Data/iftar_spec.docx'  # ملف خطاب المراد ارسالة الخطاب نسخة منه
    outLetter_path = basePath + "\\outLetter"
    outpdf_path = basePath + "\\outpdf"
    encrypt_pdf_path = basePath + "\\iftar_spec"

    if not os.path.exists(outLetter_path):
        print('Creating output folder...')
        os.makedirs(outLetter_path)
        print(outLetter_path, 'created.')
    else:
        print(outLetter_path, 'already exists.\n')

    if not os.path.exists(outpdf_path):
        print('Creating output folder...')
        os.makedirs(outpdf_path)
        print(outpdf_path, 'created.')
    else:
        print(outpdf_path, 'already exists.\n')

    if not os.path.exists(encrypt_pdf_path):
        print('Creating output folder...')
        os.makedirs(encrypt_pdf_path)
        print(encrypt_pdf_path, 'created.')
    else:
        print(encrypt_pdf_path, 'already exists.\n')

    wb = openpyxl.load_workbook(Excl_file)
    sheet = wb['Data']

    doc = docx.Document(Word_file)
    NumRecord = sheet.max_row  # تحديد عدد الاسطر في الجدول
    # NumRecord -=97
    print('Number of Records : ', NumRecord)
    lettercount = 0
    r = 3
    while (r < (NumRecord - 1)):
        lettercount += 1
        county = str(sheet.cell(row=r, column=4).value)
        collect = str(sheet.cell(row=r, column=6).value)
        nextcollect = str(sheet.cell(row=r + 1, column=6).value)
        donorName = sheet.cell(row=r, column=10).value
        centerName = str(sheet.cell(row=r, column=2).value)
        assosiation = sheet.cell(row=r, column=5)
        receiverName = sheet.cell(row=r, column=9)
        numMeal = sheet.cell(row=r, column=13)
        Email = sheet.cell(row=r, column=7)
        outNum = sheet.cell(row=r, column=15)
        outDate = sheet.cell(row=r, column=16)
        salesName = sheet.cell(row=r, column=11)
        salesMobile = sheet.cell(row=r, column=12)
        flag = sheet.cell(row=r, column=17)

        # to = Email.value
        doc_file = os.path.join(outLetter_path, r"{}.docx".format(outNum.value))
        out_file = os.path.join(outpdf_path, r"{}.pdf".format(outNum.value))
        encrypted_file = os.path.join(encrypt_pdf_path, r"{}.pdf".format(outNum.value))
        temppdf = os.path.join(encrypt_pdf_path, r"temp.pdf")
        dist = assosiation.value + '-' + county
        doc.paragraphs[0].runs[8].text = str(outNum.value)
        doc.paragraphs[1].runs[1].text = str(outDate.value)
        doc.paragraphs[3].runs[1].text = dist
        doc.paragraphs[5].runs[6].text = receiverName.value
        doc.paragraphs[8].runs[5].text = str(salesMobile.value)
        doc.paragraphs[8].runs[1].text = salesName.value
        # doc.paragraphs[3].runs[2].text = ' '*(68-int(2.3+len(dist)))

        recordOrder = 1
        # print(NumRecord, r)

        tb = doc.tables[0]
        total = numMeal.value

        clist = [[str(recordOrder), centerName, donorName, str(numMeal.value)]]

        while collect == nextcollect:
            r += 1
            recordOrder += 1
            donorName = sheet.cell(row=r, column=10).value
            centerName = sheet.cell(row=r, column=2).value
            numMeal = sheet.cell(row=r, column=13)

            collect = sheet.cell(row=r, column=6).value
            nextcollect = sheet.cell(row=r + 1, column=6).value

            clist.append([str(recordOrder), centerName, donorName, str(numMeal.value)])
            total += numMeal.value
        else:
            pass
        # print('clist lingth: ',len(clist))
        for item in clist:
            cells = tb.add_row().cells
            cells[0].text = item[0]
            cells[1].text = item[1]
            cells[2].text = item[2]
            cells[3].text = item[3]
        # print('total is : ',total)
        doc.paragraphs[5].runs[8].text = str(total)
        section = doc.sections
        # print(len(section))
        sec = section[1]
        sec.top_margin = Inches(2.5)

        doc.save(doc_file)
        # print('record after file is : ', recordOrder)
        rowtable = len(tb.rows)
        # print('Number of row table  is : ', rowtable)
        # جزء البرنامج الذي يقوم بتحويل الخطاب من ورد الى بي دي اف
        wdFormatPDF = 17
        in_file = doc_file
        # print('maher')
        word = comtypes.client.CreateObject('Word.Application')
        docfinal = word.Documents.Open(in_file)
        docfinal.SaveAs(temppdf, FileFormat=wdFormatPDF)
        # print(temppdf)

        # print('finsh')
        docfinal.Close()
        word.Quit()

        clist.clear()
        # os.system('taskkill /F /IM WINWORD.EXE')
        for i in range(rowtable - 1, 0, -1):
            row = tb.rows[i]
            # print('row is : ', row)
            remove_row(tb, row)
        # print("Number of rows after delet  :  ",len(tb.rows))

        # **************************************

        #    التشفير باستخدام ادوات pdftk
        #    %%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%
        cmd = 'pdftk {} cat 1-r2 output {} '.format(temppdf, out_file)
        os.system(cmd)
        os.system('del /F {}'.format(temppdf))
        cmd = 'pdftk "{0}" input_pw "{1}" background background.pdf output "{2}" ' \
              'encrypt_128bit allow "printing"'.format(out_file, owner_password, encrypted_file)
        os.system(cmd)
        # subprocess.call(cmd)
        # %%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%
        # جزء البرنامج الذي يقوم بارسال بريد الكتروني لكل شركة
        # fo Email send information
        sender = "social@wamy.org"
        subject = 'تعميد الجهات المشرفة علي إفطار الصائم 1440هـ'
        message_text_plain = ""
        attached_file = encrypted_file
        message_text_html = """<blockquote>
        <p><b><div align="center" ><h2>   الاخوة الأفاضل  : {}</h2><br>
         آمل ان يصلكم خطابنا هذا وأنتم في صحة طيبة وعلى خير ما يحب ربنا ويرضى   <br>  
           مرفق لكم تعميدكم بعدد وجبات  إفطار الصائم لعام 1440 هـ   <br>

           نأمل التواصل مع ( <font size="4" color="blue">{}</font>) لتنفيذ البرنامج  <br><br>

         <b><h2>الندوة العالمية للشباب الاسلامي<br>
           إدارة الشؤون الإجتماعية<br> 
                  </b></h2>
        </div></blockquote>""".format(assosiation.value, receiverName.value)
        if outNum.value != None:
            # create_message_and_send(sender, to, subject, message_text_plain, message_text_html, attached_file)
            flag.alignment = Alignment(horizontal='center', vertical='center')
            flag.font = Font(color="FF0000")
            flag.value = 1
            # window.done(to, assosiation.value, outNum.value)
            time.sleep(0.2)

        else:
            # print("لا يوجد بريد لارساله ")
            flag.value = 0
            time.sleep(0.2)

        r += 1
        window.textEdit.append('جاري العمل على السطر     {}'.format(r))

    wb.save(Excl_file)
    totalPdf(encrypt_pdf_path, basePath, sheet.cell(row=5, column=15).value[:4])
    print("done")
    window.textEdit.append('تم انشاء     {} خطاب'.format(lettercount))


def createQuestionnaire():
    window.textEdit.clear()
    basePath = os.getcwd()
    print('this is base path : ', os.getcwd())
    owner_password = "Wamy@12379"
    Word_file = 'word_Data/questionnaire.docx'  # ملف خطاب المراد ارسالة الخطاب نسخة منه
    outLetter_path = basePath + "\\outLetter"
    outpdf_path = basePath + "\\outpdf"
    encrypt_pdf_path = basePath + "\\question"

    if not os.path.exists(outLetter_path):
        print('Creating output folder...')
        os.makedirs(outLetter_path)
        print(outLetter_path, 'created.')
    else:
        print(outLetter_path, 'already exists.\n')

    if not os.path.exists(outpdf_path):
        print('Creating output folder...')
        os.makedirs(outpdf_path)
        print(outpdf_path, 'created.')
    else:
        print(outpdf_path, 'already exists.\n')

    if not os.path.exists(encrypt_pdf_path):
        print('Creating output folder...')
        os.makedirs(encrypt_pdf_path)
        print(encrypt_pdf_path, 'created.')
    else:
        print(encrypt_pdf_path, 'already exists.\n')
    doc = docx.Document(Word_file)
    doc_file = os.path.join(outLetter_path, r"questionnaire.docx")
    out_file = os.path.join(outpdf_path, r"questionnaire.pdf")
    encrypted_file = os.path.join(encrypt_pdf_path, r"questionnaire.pdf")
    doc.save(doc_file)
    # جزء البرنامج الذي يقوم بتحويل الخطاب من ورد الى بي دي اف
    wdFormatPDF = 17
    word = comtypes.client.CreateObject('Word.Application')
    docfinal = word.Documents.Open(doc_file)
    docfinal.SaveAs(out_file, FileFormat=wdFormatPDF)
    docfinal.Close()
    word.Quit()

    # **************************************
    #    التشفير باستخدام ادوات pdftk
    #    %%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%

    cmd = 'pdftk "{}" input_pw "{}" background background.pdf output "{}" encrypt_128bit allow "printing"'.format(
        out_file, owner_password,
        encrypted_file)

    os.system(cmd)
    window.textEdit.append('Creating Questionnaire Done \n you can find it in ({})folder'.format(encrypt_pdf_path))


def sendََQuestionnaire():
    window.textEdit.clear()
    basePath = os.getcwd()
    print('this is base path : ', os.getcwd())
    Excl_file = "excel_Data/questionnaire.xlsx"  # ملف اكسل يحتوي على معلومات كل الجهات المراد ارسالها
    Word_file = 'word_Data/questionnaire.docx'  # ملف خطاب المراد ارسالة الخطاب نسخة منه
    outLetter_path = basePath + "\\outLetter"
    outpdf_path = basePath + "\\outpdf"
    encrypt_pdf_path = basePath + "\\question"
    encrypted_file = os.path.join(encrypt_pdf_path, r"questionnaire.pdf")

    if not os.path.exists(encrypt_pdf_path):
        print('No folder contain questionnaire.pdf')
        window.textEdit.append("No folder contain questionnaire.pdf")
    elif not os.path.exists(encrypted_file):
        window.textEdit.append('no questionnaire.pdf file exist\n')
        print(encrypted_file, 'not exist')
    else:
        window.textEdit.append("questionnaire.pdf is existing we'll send this file to all distenaion\n")

    wb = openpyxl.load_workbook(Excl_file)
    sheet = wb['Data']

    doc = docx.Document(Word_file)
    NumRecord = sheet.max_row  # تحديد عدد الاسطر في الجدول
    print('Number of Records : ', NumRecord)
    for r in range(5, NumRecord + 1):
        Email = sheet.cell(row=r, column=2)
        status = sheet.cell(row=r, column=3)
        to = Email.value
        # **************************************
        # جزء البرنامج الذي يقوم بارسال بريد الكتروني لكل شركة
        # for Email send information
        sender = "social@wamy.org"
        subject = 'استبيان إفطار الصائم 1440هـ'
        message_text_plain = ""
        attached_file = encrypted_file
        message_text_html = """<blockquote>
        <p><b><div align="center" ><h2>   الاخوة الأفاضل / مديري مكاتب الندوة العالمية للشباب الاسلامي و روؤساء الجمعيات  </h2><br>
         آمل ان يصلكم خطابنا هذا وأنتم في صحة طيبة وعلى خير ما يحب ربنا ويرضى   <br>  
           مرفق لكم خطاب طلب استبيان إفطار الصائم لعام 1440 هـ نأمل منكم التوجه للرابط التالي لمىء الاستبيان  <br>
           <br> https://goo.gl/QxeCLM
           
         <b><h2>الندوة العالمية للشباب الاسلامي<br>
           إدارة الشؤون الإجتماعية<br> 
                  </b></h2>
        </div></blockquote>"""
        if to != None:
            print("now send Email to : {}".format(to))
            create_message_and_send(sender, to, subject, message_text_plain, message_text_html, attached_file)
            status.alignment = Alignment(horizontal='center', vertical='center')
            status.font = Font(color="FF0000")
            status.value = " تم الارسال بنجاح"
            window.textEdit.append('Record  ({0})Email send to the    ({1})       successfully'.format(r, to))
            time.sleep(0.2)
        else:
            print("لا يوجد بريد لارساله ")
            window.textEdit.append('No Email Address for    {}   record'.format(r))
            status.value = "لا يوجد عنوان بريد "
            time.sleep(0.2)

    wb.save(Excl_file)
    print("done Email send to all distenations ")


def aytamCreate():
    window.textEdit.clear()
    # print('aytam creation letter')
    basePath = os.getcwd()
    print('this is base path : ', os.getcwd())
    owner_password = "Wamy@12379"
    Excl_file = "excel_Data/aytam.xlsx"  # ملف اكسل يحتوي على معلومات كل الجهات المراد ارسالها
    Word_file = 'word_Data/aytam.docx'  # ملف خطاب المراد ارسالة الخطاب نسخة منه
    outLetter_path = basePath + "\\outLetter"
    outpdf_path = basePath + "\\outpdf"
    encrypt_pdf_path = basePath + "\\aytam"
    if not os.path.exists(outLetter_path):
        print('Creating output folder...')
        os.makedirs(outLetter_path)
        print(outLetter_path, 'created.')
    else:
        print(outLetter_path, 'already exists.\n')

    if not os.path.exists(outpdf_path):
        print('Creating output folder...')
        os.makedirs(outpdf_path)
        print(outpdf_path, 'created.')
    else:
        print(outpdf_path, 'already exists.\n')

    if not os.path.exists(encrypt_pdf_path):
        print('Creating output folder...')
        os.makedirs(encrypt_pdf_path)
        print(encrypt_pdf_path, 'created.')
    else:
        print(encrypt_pdf_path, 'already exists.\n')
    wb = openpyxl.load_workbook(Excl_file)
    sheet = wb['Data']
    doc = docx.Document(Word_file)
    NumRecord = sheet.max_row  # تحديد عدد الاسطر في الجدول
    print('Number of Records : ', NumRecord)
    for r in range(5, NumRecord + 1):

        assosiation = sheet.cell(row=r, column=2)
        receiverName = sheet.cell(row=r, column=9)
        numMeal = sheet.cell(row=r, column=12)
        outNum = sheet.cell(row=r, column=13)
        outDate = sheet.cell(row=r, column=14)
        salesName = sheet.cell(row=r, column=10)
        salesMobile = sheet.cell(row=r, column=11)
        flag = sheet.cell(row=r, column=15)
        status = sheet.cell(row=r, column=16)
        if assosiation.value != None:
            doc.paragraphs[0].runs[3].text = str(outNum.value)
            doc.paragraphs[1].runs[3].text = str(outDate.value)
            doc.paragraphs[6].runs[1].text = assosiation.value
            doc.paragraphs[10].runs[1].text = receiverName.value
            doc.paragraphs[11].runs[3].text = str(salesMobile.value)
            doc.paragraphs[11].runs[1].text = salesName.value
            doc.paragraphs[10].runs[3].text = str(numMeal.value)
            doc_file = os.path.join(outLetter_path, r"{}.docx".format(outNum.value))
            out_file = os.path.join(outpdf_path, r"{}.pdf".format(outNum.value))
            encrypted_file = os.path.join(encrypt_pdf_path, r"{}.pdf".format(outNum.value))
            doc.save(doc_file)
            print('تم توليد الخطاب رقم : ', outNum.value)
            # جزء البرنامج الذي يقوم بتحويل الخطاب من ورد الى بي دي اف
            wdFormatPDF = 17
            in_file = doc_file
            word = comtypes.client.CreateObject('Word.Application')
            docfinal = word.Documents.Open(in_file)
            docfinal.SaveAs(out_file, FileFormat=wdFormatPDF)
            docfinal.Close()
            word.Quit()

            #    التشفير باستخدام ادوات pdftk
            cmd = 'pdftk "{}" input_pw "{}" background background.pdf output "{}" encrypt_128bit allow "printing"'.format(
                out_file, owner_password,
                encrypted_file)
            os.system(cmd)
            if outNum.value != None:
                status.alignment = Alignment(horizontal='center', vertical='center')
                status.font = Font(color="FF0000")
                flag.value = 1
            else:
                print("لا يوجد اسم صادر  ")
                flag.value = 0
        else:
            print('لا يوجد اسم جهة')
            break
    wb.save(Excl_file)
    # os.system('taskkill /F /IM WINWORD.EXE')
    totalPdf(encrypt_pdf_path, basePath, sheet.cell(row=5, column=13).value[:4])
    print("done")
    window.textEdit.append('تم انشاء خطابات الايتام')


def aytamSend():
    window.textEdit.clear()
    print('aytam send letter')
    basePath = os.getcwd()
    Excl_file = "excel_Data/aytam.xlsx"  # ملف اكسل يحتوي على معلومات كل الجهات المراد ارسالها
    encrypt_pdf_path = basePath + "\\aytam"
    wb = openpyxl.load_workbook(Excl_file)
    sheet = wb['Data']
    NumRecord = sheet.max_row  # تحديد عدد الاسطر في الجدول
    print('Number of Records : ', NumRecord)
    r = 5
    while (r < (NumRecord + 1)):
        assosiation = sheet.cell(row=r, column=2)
        receiverName = sheet.cell(row=r, column=9)
        Email = sheet.cell(row=r, column=8)
        outNum = sheet.cell(row=r, column=13)
        flag = sheet.cell(row=r, column=15).value
        status = sheet.cell(row=r, column=16)

        # to = 'r_maher@wamy.org'
        # Email_To_Send =[Email]
        # to = 'amedshel@wamy.org'
        to = Email.value

        encrypted_file = os.path.join(encrypt_pdf_path, r"{}.pdf".format(outNum.value))
        print(encrypted_file)
        # جزء البرنامج الذي يقوم بارسال بريد الكتروني لكل شركة
        # fo Email send information

        sender = "social@wamy.org"
        subject = 'برنامج مستحقات الأيتام لعام 1440هـ'
        message_text_plain = ""
        attached_file = encrypted_file
        message_text_html = """<blockquote>
                       <p><b><div align="center" ><h2>   الاخوة الأفاضل  : {}</h2><br>
                        آمل ان يصلكم خطابنا هذا وأنتم في صحة طيبة وعلى خير ما يحب ربنا ويرضى   <br>  
                          مرفق لكم تعميدكم تنفيذ برنامج الأيتام الخاص بكم لعام 1440 هـ   <br>

                          نأمل التواصل مع ( <font size="4" color="blue">{}</font>) لتنفيذ البرنامج  <br><br>

                        <b><h2>الندوة العالمية للشباب الاسلامي<br>
                          إدارة الشؤون الإجتماعية<br> 
                                 </b></h2>
                       </div></blockquote>""".format(assosiation.value, receiverName.value)
        if flag == 1:
            print('flag = 1')
            if to != None:
                print("now send Email to : {}".format(to))
                create_message_and_send(sender, to, subject, message_text_plain, message_text_html, attached_file)
                status.alignment = Alignment(horizontal='center', vertical='center')
                status.font = Font(color="FF0000")
                status.value = " تم الارسال بنجاح"
                window.textEdit.append('recored ({0})   Email Send to ({1}) successfully '.format(r, to))
                # window.done(to, assosiation.value, outNum.value)
                time.sleep(0.2)
            else:
                print("لا يوجد بريد لارساله ")
                status.value = "لا يوجد عنوان بريد "
                window.textEdit.append('No Email address in recored ({})'.format(r))
                time.sleep(0.2)
        else:
            print('Flag = 0')
            window.textEdit.append('No letter for record ({})'.format(r))
        r += 1
    wb.save(Excl_file)
    window.textEdit.append('Email send to all distenation successfully ')
    print('done......')


def adahiCreate():
    pdfListName = []
    basePath = os.getcwd()
    print('this is base path : ', os.getcwd())
    owner_password = "Wamy@12379"
    Excl_file = "excel_Data/adahi.xlsx"  # ملف اكسل يحتوي على معلومات كل الجهات المراد ارسالها
    Word_file = 'word_Data/adahi.docx'  # ملف خطاب المراد ارسالة الخطاب نسخة منه
    outLetter_path = basePath + "\\outLetter"
    outpdf_path = basePath + "\\outpdf"
    encrypt_pdf_path = basePath + "\\adahi"
    if not os.path.exists(outLetter_path):
        print('Creating output folder...')
        os.makedirs(outLetter_path)
        print(outLetter_path, 'created.')
    else:
        print(outLetter_path, 'already exists.\n')

    if not os.path.exists(outpdf_path):
        print('Creating output folder...')
        os.makedirs(outpdf_path)
        print(outpdf_path, 'created.')
    else:
        print(outpdf_path, 'already exists.\n')

    if not os.path.exists(encrypt_pdf_path):
        print('Creating output folder...')
        os.makedirs(encrypt_pdf_path)
        print(encrypt_pdf_path, 'created.')
    else:
        print(encrypt_pdf_path, 'already exists.\n')

    wb = openpyxl.load_workbook(Excl_file)
    sheet = wb['Data']

    doc = docx.Document(Word_file)
    NumRecord = sheet.max_row  # تحديد عدد الاسطر في الجدول
    print('Number of Records : ', NumRecord)
    # totalpdfpages = os.path.join(encrypt_pdf_path,r"Adahi.pdf")

    for r in range(5, NumRecord):
        assosiation = sheet.cell(row=r, column=2)
        office = sheet.cell(row=r, column=3).value
        country = sheet.cell(row=r, column=4).value
        receiverName = sheet.cell(row=r, column=9)
        numMeal = sheet.cell(row=r, column=12)
        Email = sheet.cell(row=r, column=8)
        outNum = sheet.cell(row=r, column=13)
        outDate = sheet.cell(row=r, column=14)
        salesName = sheet.cell(row=r, column=10)
        salesMobile = sheet.cell(row=r, column=11)
        flag = sheet.cell(row=r, column=15)
        status = sheet.cell(row=r, column=16)
        copy = sheet.cell(row=r, column=17).value
        to = Email.value
        print('تم توليد الخطاب رقم : ', outNum.value)
        doc.paragraphs[0].runs[3].text = str(outNum.value)
        doc.paragraphs[1].runs[3].text = str(outDate.value)
        doc.paragraphs[6].runs[1].text = assosiation.value + '-' + country
        doc.paragraphs[10].runs[7].text = receiverName.value
        # اضافة صفر واحد على رقم الجوال المقروء من جدول اكسل اذا كانت الصيغة المطلوبة محلية وصفرين اذا كانت الصيغة دولية
        doc.paragraphs[11].runs[3].text = str(salesMobile.value)
        doc.paragraphs[11].runs[1].text = salesName.value
        doc.paragraphs[10].runs[9].text = str(numMeal.value)

        if copy == 1:
            doc.paragraphs[23].runs[0].text = '-'
            doc.paragraphs[23].runs[1].text = 'صورة مع التحية لمدير مكتب {} '.format(country)
        else:
            doc.paragraphs[23].runs[1].text = ''
            doc.paragraphs[23].runs[0].text = ''

        doc_file = os.path.join(outLetter_path, r"{}.docx".format(outNum.value))
        out_file = os.path.join(outpdf_path, r"{}.pdf".format(outNum.value))
        encrypted_file = os.path.join(encrypt_pdf_path, r"{}.pdf".format(outNum.value))

        doc.save(doc_file)
        # جزء البرنامج الذي يقوم بتحويل الخطاب من ورد الى بي دي اف
        wdFormatPDF = 17
        in_file = doc_file

        word = comtypes.client.CreateObject('Word.Application')
        docfinal = word.Documents.Open(in_file)
        docfinal.SaveAs(out_file, FileFormat=wdFormatPDF)
        docfinal.Close()
        word.Quit()

        # **************************************
        #    التشفير باستخدام ادوات pdftk
        #    %%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%
        # print(out_file)

        cmd = 'pdftk "{}" input_pw "{}" background background.pdf output "{}" encrypt_128bit allow "printing"'.format(
            out_file, owner_password,
            encrypted_file)
        pdfListName.append(out_file)

        os.system(cmd)

        if outNum.value != None:
            status.alignment = Alignment(horizontal='center', vertical='center')
            status.font = Font(color="FF0000")
            flag.value = 1
            window.textEdit.append('تم توليد الخطاب رقم :   {} '.format(outNum.value))

        else:
            print("لا يوجد اسم صادر  ")
            flag.value = 0

    for item in pdfListName:
        print('item : ', item)
    wb.save(Excl_file)
    totalPdf(encrypt_pdf_path, basePath, sheet.cell(row=5, column=13).value[:4])

    # os.chdir(os.path.join(basePath, r"adahi"))
    #
    # cmd = 'pdftk *.pdf cat output Total_adahi_{}.pdf'.format(sheet.cell(row=5, column=13).value[:4])
    # os.system(cmd)
    # os.system('taskkill /F /IM WINWORD.EXE')
    # os.chdir(basePath)
    print("done")


def adahiSend():
    window.textEdit.clear()
    print('sending email spec. Letter ')
    basePath = os.getcwd()
    Excl_file = "excel_Data/adahi.xlsx"  # ملف اكسل يحتوي على معلومات كل الجهات المراد ارسالها
    encrypt_pdf_path = basePath + "\\adahi"
    wb = openpyxl.load_workbook(Excl_file)
    sheet = wb['Data']
    NumRecord = sheet.max_row  # تحديد عدد الاسطر في الجدول
    print('Number of Records : ', NumRecord)
    r = 5
    # r =    102
    while (r < (NumRecord)):
        assosiation = sheet.cell(row=r, column=2)
        receiverName = sheet.cell(row=r, column=9)
        Email = sheet.cell(row=r, column=8)
        outNum = sheet.cell(row=r, column=13)
        flag = sheet.cell(row=r, column=15).value
        status = sheet.cell(row=r, column=16)

        # to = 'r_maher@wamy.org'
        # Email_To_Send =[Email]
        # to = 'amedshel@wamy.org'
        to = Email.value
        print('test')

        encrypted_file = os.path.join(encrypt_pdf_path, r"{}.pdf".format(outNum.value))
        print(encrypted_file)
        # جزء البرنامج الذي يقوم بارسال بريد الكتروني لكل شركة
        # fo Email send information
        sender = "social@wamy.org"
        subject = 'تعميد الجهات المشرفة علي إفطار الصائم 1440هـ'
        message_text_plain = ""
        attached_file = encrypted_file
        message_text_html = """<blockquote>
                       <p><b><div align="center" ><h2>   الاخوة الأفاضل  : {}</h2><br>
                        آمل ان يصلكم خطابنا هذا وأنتم في صحة طيبة وعلى خير ما يحب ربنا ويرضى   <br>  
                          مرفق لكم تعميدكم بعدد الأضاحي  لعام 1440 هـ   <br>

                          نأمل التواصل مع ( <font size="4" color="blue">{}</font>) لتنفيذ البرنامج  <br><br>

                        <b><h2>الندوة العالمية للشباب الاسلامي<br>
                          إدارة الشؤون الإجتماعية<br> 
                                 </b></h2>
                       </div></blockquote>""".format(assosiation.value, receiverName.value)
        if flag == 1:
            print('flag = 1')
            if to != None:
                print("now send Email to : {}".format(to))
                create_message_and_send(sender, to, subject, message_text_plain, message_text_html, attached_file)
                status.alignment = Alignment(horizontal='center', vertical='center')
                status.font = Font(color="FF0000")
                status.value = " تم الارسال بنجاح"
                window.done(to, assosiation.value, outNum.value)
                time.sleep(0.2)
            else:
                print("لا يوجد بريد لارساله ")
                status.value = "لا يوجد عنوان بريد "
                time.sleep(0.2)
        else:
            print('Flag = 0')
        r += 1
    wb.save(Excl_file)
    print('done......')


class Window(QMainWindow):
    def __init__(self):
        super().__init__()

        self.title = "ارسال الخطابات الآلي المشفر"
        self.top = 100
        self.left = 100
        self.width = 800
        self.height = 700
        self.setWindowIcon(QtGui.QIcon("pdf1.png"))

        self.InitWindow()

    def InitWindow(self):
        exitAct = QAction(QIcon('exit.png'), 'Exit', self)
        exitAct.setShortcut('ctrl+Q')
        exitAct.triggered.connect(self.CloseApp)

        copyAct = QAction(QIcon('copy.png'), 'Copy', self)
        copyAct.setShortcut('ctrl+C')

        pasteAct = QAction(QIcon('paste.png'), 'Paste', self)
        pasteAct.setShortcut('ctrl+V')

        deleteAct = QAction(QIcon('delete.png'), 'Delete', self)
        deleteAct.setShortcut('ctrl+D')

        saveAct = QAction(QIcon('save.png'), 'Save', self)
        saveAct.setShortcut('ctrl+S')

        self.toolbar = self.addToolBar('ToolBar')

        self.toolbar.addAction(exitAct)
        self.toolbar.addAction(copyAct)
        self.toolbar.addAction(pasteAct)
        self.toolbar.addAction(deleteAct)
        self.toolbar.addAction(saveAct)

        x = 100
        y = 50

        self.button9 = QPushButton("ارسال الخطابات المخصصة", self)
        self.button9.setGeometry(x, y, 300, 50)
        self.button9.clicked.connect(sendEmail_spec)

        self.button8 = QPushButton("انشاءالخطابات المخصصة", self)
        self.button8.setGeometry(x + 300, y, 300, 50)
        self.button8.clicked.connect(create_Spec)

        self.button6 = QPushButton("انشاء الخطابات غير المخصص", self)
        self.button6.setGeometry(x + 300, y + 50, 300, 50)
        self.button6.clicked.connect(createLetter)

        self.button7 = QPushButton("ارسال الخطابات غير المخصصة", self)
        self.button7.setGeometry(x, y + 50, 300, 50)
        self.button7.clicked.connect(sendLetter)

        self.button10 = QPushButton("ارسال تعميم استبيان", self)
        self.button10.setGeometry(x, y + 150, 300, 50)
        self.button10.clicked.connect(sendََQuestionnaire)

        self.button5 = QPushButton("انشاء تعميم استبيان", self)
        self.button5.setGeometry(x + 300, y + 150, 300, 50)
        self.button5.clicked.connect(createQuestionnaire)

        self.button6 = QPushButton("انشاء خطابات الايتام", self)
        self.button6.setGeometry(x + 300, y + 100, 300, 50)
        self.button6.clicked.connect(aytamCreate)

        self.button7 = QPushButton("ارسال خطابات الايتام", self)
        self.button7.setGeometry(x, y + 100, 300, 50)
        self.button7.clicked.connect(aytamSend)

        self.button6 = QPushButton("انشاء خطابات الاضاحي", self)
        self.button6.setGeometry(x + 300, y + 200, 300, 50)
        self.button6.clicked.connect(adahiCreate)

        self.button7 = QPushButton("ارسال خطابات الأضاحي", self)
        self.button7.setGeometry(x, y + 200, 300, 50)
        self.button7.clicked.connect(adahiSend)

        self.button1 = QPushButton("Open Font Dialog", self)
        self.button1.setGeometry(x, y + 300, 200, 50)
        self.button1.clicked.connect(self.createFontDialog)

        self.button2 = QPushButton("Open Color Dialog", self)
        self.button2.setGeometry(x + 200, y + 300, 200, 50)
        self.button2.clicked.connect(self.creatColorDialog)

        self.button3 = QPushButton("Print ", self)
        self.button3.setGeometry(x + 400, y + 300, 100, 50)
        self.button3.clicked.connect(self.creatPrintDialog)

        self.button4 = QPushButton("Print Preview", self)
        self.button4.setGeometry(x + 500, y + 300, 100, 50)
        self.button4.clicked.connect(self.creatPrintPreViewDialog)

        self.textEdit = QTextEdit(self)
        self.textEdit.setGeometry(x, y + 350, 600, 200)

        self.setWindowTitle(self.title)
        self.setGeometry(self.top, self.left, self.width, self.height)
        self.show()

    def CloseApp(self):
        self.close()

    def createFontDialog(self):
        font, ok = QFontDialog.getFont()

        if ok:
            self.textEdit.setFont(font)

    def creatColorDialog(self):
        color = QColorDialog.getColor()
        self.textEdit.setTextColor(color)

    def creatPrintDialog(self):
        printer = QPrinter(QPrinter.HighResolution)
        dialog = QPrintDialog(printer, self)

        if dialog.exec_() == QPrintDialog.Accepted:
            self.textEdit.print_(printer)

    def creatPrintPreViewDialog(self):
        printer = QPrinter(QPrinter.HighResolution)
        previewDialog = QPrintPreviewDialog(printer, self)
        previewDialog.paintRequested.connect(self.printPreview)
        previewDialog.exec_()

    def printPreview(self, printer):
        self.textEdit.print_(printer)

    def done(self, Emeil, assosiation, outNum):

        self.textEdit.append(
            "تم الارسال الى{1} {3} الجهة : {0}{4} رقم الصادر {2} بنجاح ".format(assosiation, Emeil, outNum,
                                                                                ' ' * (35 - len(Emeil)),
                                                                                ' ' * (30 - len(assosiation))))


App = QApplication(sys.argv)
window = Window()
sys.exit(App.exec())
